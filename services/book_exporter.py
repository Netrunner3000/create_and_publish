"""
Manuscript export for Create & Publish's Manuscript (author) agent.
Splits a plain-text draft into chapters and renders it as EPUB, DOCX, or PDF.

No external system dependencies (no LibreOffice/Calibre required):
- EPUB via EbookLib (already a project dependency).
- DOCX via python-docx.
- PDF via reportlab (pure Python, generated directly — not a DOCX conversion).
"""

from __future__ import annotations
import html
import re
import uuid
from pathlib import Path

CHAPTER_HEADING_RE = re.compile(
    r"^\s*(?:#{1,3}\s*)?(chapter\s+\S+|part\s+\S+|prologue|epilogue)\b.*$",
    re.IGNORECASE,
)


def split_into_chapters(text: str) -> list[tuple[str, str]]:
    """
    Split manuscript text into (heading, body) pairs based on heading lines
    (Chapter N / Part N / Prologue / Epilogue, optionally markdown-prefixed).
    Text before the first detected heading becomes an unlabeled opening chapter.
    Falls back to a single unlabeled chapter if no headings are found.
    """
    lines = text.splitlines()
    chapters: list[tuple[str, list[str]]] = []
    current_heading = None
    current_body: list[str] = []

    for line in lines:
        if CHAPTER_HEADING_RE.match(line.strip()):
            if current_heading is not None or current_body:
                chapters.append((current_heading or "", current_body))
            current_heading = line.strip().lstrip("#").strip()
            current_body = []
        else:
            current_body.append(line)

    if current_heading is not None or current_body:
        chapters.append((current_heading or "", current_body))

    result = [(h, "\n".join(b).strip()) for h, b in chapters if h or "\n".join(b).strip()]
    return result or [("", text.strip())]


def find_chapter_offsets(text: str) -> list[int]:
    """Character offsets in `text` where each detected chapter heading line begins,
    in document order. Used to jump an editor cursor to a chapter (see split_into_chapters
    for the corresponding (heading, body) pairs — headed chapters map 1:1 to these offsets,
    in order; an unlabeled leading chapter has no corresponding entry here)."""
    offsets = []
    pos = 0
    for line in text.splitlines(keepends=True):
        if CHAPTER_HEADING_RE.match(line.strip()):
            offsets.append(pos)
        pos += len(line)
    return offsets


def _escape(text: str) -> str:
    return html.escape(text, quote=False)


def export_epub(title: str, author_name: str, chapters: list[tuple[str, str]], output_path: Path) -> Path:
    from ebooklib import epub

    book = epub.EpubBook()
    book.set_identifier(str(uuid.uuid4()))
    book.set_title(title)
    book.set_language("en")
    book.add_author(author_name)

    epub_chapters = []
    for i, (heading, body) in enumerate(chapters, start=1):
        heading_text = heading or (title if len(chapters) == 1 else f"Chapter {i}")
        paragraphs = "".join(f"<p>{_escape(p)}</p>" for p in body.split("\n\n") if p.strip())
        c = epub.EpubHtml(title=heading_text, file_name=f"chap_{i:02d}.xhtml", lang="en")
        c.content = f"<h1>{_escape(heading_text)}</h1>{paragraphs}"
        book.add_item(c)
        epub_chapters.append(c)

    book.toc = tuple(epub_chapters)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav"] + epub_chapters

    output_path.parent.mkdir(parents=True, exist_ok=True)
    epub.write_epub(str(output_path), book)
    return output_path


def export_docx(title: str, author_name: str, chapters: list[tuple[str, str]], output_path: Path) -> Path:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(28)

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_p.add_run(author_name).font.size = Pt(14)

    doc.add_page_break()

    for i, (heading, body) in enumerate(chapters, start=1):
        heading_text = heading or (title if len(chapters) == 1 else f"Chapter {i}")
        doc.add_heading(heading_text, level=1)
        for para in body.split("\n\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)
        if i < len(chapters):
            doc.add_page_break()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def export_pdf(title: str, author_name: str, chapters: list[tuple[str, str]], output_path: Path) -> Path:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(output_path), pagesize=letter,
        topMargin=1 * inch, bottomMargin=1 * inch,
        leftMargin=1.1 * inch, rightMargin=1.1 * inch,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("BookTitle", parent=styles["Title"], fontSize=28, alignment=TA_CENTER, spaceAfter=12)
    author_style = ParagraphStyle("BookAuthor", parent=styles["Normal"], fontSize=14, alignment=TA_CENTER)
    chapter_style = ParagraphStyle("ChapterHeading", parent=styles["Heading1"], spaceBefore=24, spaceAfter=12)
    body_style = ParagraphStyle("BookBody", parent=styles["Normal"], fontSize=11, leading=16, spaceAfter=10)

    story = [
        Spacer(1, 2 * inch),
        Paragraph(_escape(title), title_style),
        Paragraph(_escape(author_name), author_style),
        PageBreak(),
    ]

    for i, (heading, body) in enumerate(chapters, start=1):
        heading_text = heading or (title if len(chapters) == 1 else f"Chapter {i}")
        story.append(Paragraph(_escape(heading_text), chapter_style))
        for para in body.split("\n\n"):
            para = para.strip()
            if para:
                story.append(Paragraph(_escape(para).replace("\n", "<br/>"), body_style))
        if i < len(chapters):
            story.append(PageBreak())

    doc.build(story)
    return output_path


EXPORTERS = {"epub": export_epub, "docx": export_docx, "pdf": export_pdf}


def export_book(text: str, title: str, author_name: str, fmt: str, output_path: Path) -> Path:
    exporter = EXPORTERS.get(fmt)
    if exporter is None:
        raise ValueError(f"Unsupported export format: {fmt!r} (use one of {list(EXPORTERS)})")
    chapters = split_into_chapters(text)
    return exporter(title or "Untitled Manuscript", author_name or "Unknown Author", chapters, output_path)
